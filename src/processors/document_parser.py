"""
This module provides functions to parse and extract data from a Word document (.docx) using the
    `python-docx` library.
It includes functions to remove paragraphs before a specific heading and to extract data between
    specific headings.
Functions:
    collect_objects_after_heading2(document: Document) -> Document:
        Removes all paragraphs in the given document that appear before the first occurrence of a
        paragraph with the style "Heading 2".
    collect_objects_between_heading3(document: Document) -> pd.DataFrame:
        Extracts data from paragraphs between "Heading 3" styled paragraphs and returns it as a
        DataFrame.
    parse_document(file_path: str) -> pd.DataFrame:
        Parses the document at the given file path, extracts relevant data, and returns it as a
        DataFrame.

Returns:
    _type_: _description_
"""

import re

import docx
import pandas as pd


def collect_objects_from_docx(document: docx.Document):
    """
    Removes all paragraphs in the given document that appear before the first
    occurrence of a paragraph with the style "Heading 2".
    Args:
        document (Document): The document object to process.
    Returns:
        Document: The modified document with paragraphs before the first
        "Heading 2" removed.
    """
    # Iterate through paragraphs
    for paragraph in document.paragraphs[:]:  # Create a copy of the list
        if paragraph.style.name == "Heading 2":
            break  # Stop processing once we find the Heading 2

        # Remove the paragraph if we haven't found Heading 2 yet
        paragraph._element.getparent().remove(paragraph._element)

    # Return the modified document
    return document  # Return the document after removing paragraphs before Heading 2


def extract_objects_from_paragraphs(paragraphs):
    """
    Extracts and collects objects between Heading 3 paragraphs in a list of paragraphs.
    This function searches through the paragraphs, identifies sections that match a specific
    pattern following a Heading 3 paragraph, and extracts relevant information to populate a
    DataFrame. The extracted information includes the title and the article text, which is
    collected until a "Back to Top" paragraph is encountered.
    Args:
        paragraphs (list): The list of paragraphs to be parsed.
    Returns:
        pd.DataFrame: A DataFrame containing the extracted data with columns "Title"
                      and "Article_Text".
    """
    pattern = (
        r"(\d+\.\d+) - (.+?): (.+) \((\d{1,2}\s(?:[A-Za-z]+))"
        r",( [\w ,-]+)?( \d+.?\w+ uvm;)?( [\w ,-]+)?\)"
    )
    df = pd.DataFrame(columns=["Title", "Article_Text"])

    for i, paragraph in enumerate(paragraphs):
        data = {}
        match = re.match(pattern, paragraph)

        if match:
            article_text = []
            for para in paragraphs[i + 2 :]:
                if para.strip() == "Back to Top":
                    break
                article_text.append(para.strip())

            data.update(
                {
                    "Title": paragraph,
                    "Article_Text": " ".join(article_text),
                }
            )

            new_data_df = pd.DataFrame([data])
            df = pd.concat([new_data_df, df], ignore_index=True)

    return df


def extract_objects_from_docx(document: docx.Document):
    """
    Extracts and collects objects between Heading 3 paragraphs in a document.
    This function searches through the paragraphs of a given document, identifies
    paragraphs that match a specific pattern following a Heading 3 paragraph, and
    extracts relevant information to populate a DataFrame. The extracted information
    includes the title and the article text, which is collected until a "Back to Top"
    paragraph is encountered.
    Args:
        document (Document): The document object containing paragraphs to be parsed.
    Returns:
        pd.DataFrame: A DataFrame containing the extracted data with columns "Title"
                      and "Article_Text".
    """
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    return extract_objects_from_paragraphs(paragraphs)


def parse_document(file_path):
    """
    Parses the document at the given file path, extracts relevant data, and appends it to the
    provided DataFrame.

    Args:
        df (pd.DataFrame): The DataFrame to which the extracted data will be appended.
        file_path (str): The path to the document file to be parsed.

    Returns:
        pd.DataFrame: The updated DataFrame with the extracted data appended.
    """
    document = docx.Document(file_path)
    document_after_heading2 = collect_objects_from_docx(document)

    df_between_heading3 = extract_objects_from_docx(document_after_heading2)
    return df_between_heading3
